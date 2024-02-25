import json
import os
from docx import Document, shared

def create_docx(books, genre):
    doc = Document()
    doc.add_heading(genre, level=1)

    for book in books:
        doc.add_heading(book['title'], level=2)
        p = doc.add_paragraph()
        p.add_run('Author: ').bold = True
        p.add_run(book['author'])
        p.add_run('\n')
        p.add_run('Publish Year: ').bold = True
        p.add_run(str(book['publish_year']))
        p.add_run('\n')
        p.add_run('Genre: ').bold = True
        p.add_run(book['genre'])

        for run in doc.paragraphs[-1].runs:
            run.font.size = shared.Pt(12)

        doc.add_paragraph()

    return doc

def main():
    # Directory to Store the output file
    output_directory = r"S:\Testing_Ground\MOD8_output"

    # open the JSON file and store it in a variable
    with open('books.json', encoding='utf-8') as f:
        data = json.load(f)

    # open a blank dictionary to store the books sorted by Genre
    books_by_genre = {}

    # Organize books by genre using lists
    for book in data['books']:
        genre = book['genre']
        if genre not in books_by_genre:
            books_by_genre[genre] = []
        books_by_genre[genre].append(book)

    # Generating the word documents for each genre
    for genre, books in books_by_genre.items():
        doc = create_docx(books, genre)
        output_path = os.path.join(output_directory, f'{genre}.docx')
        doc.save(output_path)

main()